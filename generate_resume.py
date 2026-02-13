import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def create_resume():
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Load Data
    profile_data = load_json('src/data/profile.json')
    history_data = load_json('src/data/history.json')
    projects_data = load_json('src/data/projects.json')
    stack_data = load_json('src/data/stack.json')

    profile = profile_data['profile']
    hero = profile_data['hero']

    # --- Header ---
    header = doc.add_heading(profile['name'], 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{profile['title']} | {profile['email']}").bold = True
    p.add_run(f"\n{profile['socials']['linkedin']} | {profile['socials']['github']}")

    # --- Summary ---
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(hero['description'])

    # --- Work Experience ---
    doc.add_heading('Work Experience', level=1)

    # Day-to-day tasks (Manual additions as requested)
    work_tasks = {
        "Backend Developer @ Transcosmos Asia Philippines": [
            "Develop and maintain robust RESTful APIs using Laravel and PHP for large-scale enterprise applications.",
            "Architect modular system components to improve code reusability and maintainability.",
            "Collaborate with frontend teams to integrate React components with backend services.",
            "Optimize database queries and system performance for high-traffic e-commerce platforms.",
            "Manage system integrations and third-party API connections (e.g., Shopify, Line Login)."
        ],
        "Systems Developer @ Synermaxx Corporation": [
            "Developed and enhanced enterprise e-commerce systems using PHP and various frameworks.",
            "Maintained legacy codebases and implemented new features based on client requirements.",
            "Collaborated with cross-functional teams to deliver end-to-end software solutions.",
            "Troubleshot and resolved critical production issues in a timely manner.",
            "Participated in code reviews and contributed to internal development standards."
        ],
        "Customer Service Representative @ Concentrix": [
            "Handled high volumes of customer inquiries and technical issues via phone and email.",
            "Maintained high customer satisfaction ratings through effective problem-solving.",
            "Collaborated with technical support teams to resolve complex customer concerns.",
            "Documented customer interactions and feedback accurately in CRM systems."
        ],
        "IT Trainee @ Fulrubell Corporation": [
            "Assisted in maintaining local area networks and server hardware.",
            "Provided basic IT support for staff members and troubleshooting workstation issues.",
            "Supported database administration tasks and data entry verification.",
            "Gained exposure to enterprise IT infrastructure and security protocols."
        ]
    }

    for job in history_data:
        if job['status'] == 'EDUCATION':
            continue
            
        role_header = doc.add_paragraph()
        role_header.add_run(f"{job['title']}").bold = True
        role_header.add_run(f" | {job['period']}")
        
        doc.add_paragraph().add_run(job['description']).italic = True
        
        tasks = work_tasks.get(job['title'], [])
        for task in tasks:
            doc.add_paragraph(task, style='List Bullet')

    # --- Skills ---
    doc.add_heading('Technical Skills', level=1)
    skills_list = ", ".join([skill['name'] for skill in stack_data])
    doc.add_paragraph(skills_list)

    # --- Key Projects ---
    doc.add_heading('Key Projects', level=1)
    
    # Featured Projects
    for project in projects_data['featured']:
        p = doc.add_paragraph()
        p.add_run(f"{project['title']} ({project['role']})").bold = True
        doc.add_paragraph(project['description'])
        for contribution in project['contributions']:
            doc.add_paragraph(contribution, style='List Bullet')

    # --- Education ---
    doc.add_heading('Education', level=1)
    for edu in history_data:
        if edu['status'] == 'EDUCATION':
            p = doc.add_paragraph()
            p.add_run(f"{edu['title']}").bold = True
            p.add_run(f" | {edu['period']}")
            doc.add_paragraph(edu['description'])

    # Save
    filename = "Jean-Louis_Mendoza_Resume.docx"
    doc.save(filename)
    print(f"Resume generated successfully: {filename}")

if __name__ == "__main__":
    create_resume()
